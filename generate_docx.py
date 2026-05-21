"""
Script tạo file Word chuẩn báo cáo từ BaoCao_Hanzii.md
Chuẩn: Times New Roman 13pt, lề 2.5cm (trên/dưới/phải), 3.5cm (trái), dãn dòng 1.5
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Đọc file markdown ──────────────────────────────────────────────────────────
with open("BaoCao_Hanzii.md", encoding="utf-8") as f:
    lines = f.readlines()

# ── Tạo document ──────────────────────────────────────────────────────────────
doc = Document()

# Lề trang: trái 3.5cm, phải/trên/dưới 2.5cm
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set font cho run ──────────────────────────────────────────────────
def set_run_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Buộc font áp dụng cho tiếng Việt (East Asian font)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)

def set_para_spacing(para, space_before=0, space_after=6, line_spacing=1.5):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'),  str(int(space_after * 20)))
    spacing.set(qn('w:line'),   str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    # Xóa spacing cũ nếu có
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    pPr.append(spacing)

def add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
             italic=False, indent_cm=0, space_before=0, space_after=6,
             line_spacing=1.5, color=None):
    """Thêm đoạn văn với inline bold/italic từ markdown **text** hoặc *text*"""
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, space_before, space_after, line_spacing)
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)

    # Parse inline bold/italic
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(.+?))'
    for m in re.finditer(pattern, text):
        if m.group(2):   # ***bold italic***
            run = p.add_run(m.group(2))
            set_run_font(run, size, bold=True, italic=True, color=color)
        elif m.group(3): # **bold**
            run = p.add_run(m.group(3))
            set_run_font(run, size, bold=True, color=color)
        elif m.group(4): # *italic*
            run = p.add_run(m.group(4))
            set_run_font(run, size, italic=True, color=color)
        elif m.group(5): # `code`
            run = p.add_run(m.group(5))
            run.font.name = "Courier New"
            run.font.size = Pt(11)
        elif m.group(6): # plain text
            run = p.add_run(m.group(6))
            set_run_font(run, size, bold=bold, italic=italic, color=color)
    return p

def add_heading(text, level):
    """Thêm heading với định dạng chuẩn báo cáo"""
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=12, space_after=6, line_spacing=1.5)
    clean = re.sub(r'\*+', '', text).strip()

    if level == 1:   # # Tiêu đề lớn — căn giữa, 14pt bold
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean.upper())
        set_run_font(run, 14, bold=True)
    elif level == 2: # ## Chương — căn trái, 13pt bold
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(clean)
        set_run_font(run, 13, bold=True)
    elif level == 3: # ### Mục — thụt 1cm, 13pt bold
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(clean)
        set_run_font(run, 13, bold=True)
    elif level == 4: # #### Tiểu mục — thụt 1.5cm, 13pt bold italic
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(clean)
        set_run_font(run, 13, bold=True, italic=True)
    return p

def add_code_block(lines_content):
    """Thêm khối code với font Courier New, nền xám nhạt"""
    p = doc.add_paragraph()
    set_para_spacing(p, 3, 3, 1.0)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    # Nền xám cho đoạn
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    for i, line_c in enumerate(lines_content):
        if i > 0:
            run = p.add_run('\n')
        run = p.add_run(line_c.rstrip('\n'))
        run.font.name = "Courier New"
        run.font.size = Pt(10)

def add_table_from_md(table_lines):
    """Tạo bảng từ markdown table lines"""
    rows_data = []
    for line_t in table_lines:
        line_t = line_t.strip()
        if re.match(r'^\|[-:| ]+\|$', line_t):
            continue  # separator row
        cells = [c.strip() for c in line_t.strip('|').split('|')]
        rows_data.append(cells)

    if not rows_data:
        return

    max_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Table Grid'

    for r_i, row_data in enumerate(rows_data):
        row_obj = table.rows[r_i]
        for c_i, cell_text in enumerate(row_data):
            if c_i >= max_cols:
                break
            cell = row_obj.cells[c_i]
            cell.text = ''
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p_cell, 2, 2, 1.15)
            # Parse inline bold
            clean_cell = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            clean_cell = re.sub(r'\*(.+?)\*',   r'\1', clean_cell)
            clean_cell = re.sub(r'`(.+?)`',     r'\1', clean_cell)
            run = p_cell.add_run(clean_cell)
            is_header = (r_i == 0)
            set_run_font(run, 12, bold=is_header)

    # Căn chiều rộng bảng
    from docx.oxml.ns import qn as _qn
    tbl = table._tbl
    tblPr = tbl.find(_qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(_qn('w:w'), '5000')
    tblW.set(_qn('w:type'), 'pct')
    tblPr.append(tblW)

# ── Phân tích và render từng dòng ─────────────────────────────────────────────
i = 0
in_code_block = False
code_lines_buf = []
in_table = False
table_lines_buf = []
skip_comment = False

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Bỏ qua chú thích HTML <!-- ... -->
    if '<!--' in stripped:
        if '-->' in stripped:
            i += 1
            continue
        skip_comment = True
        i += 1
        continue
    if skip_comment:
        if '-->' in stripped:
            skip_comment = False
        i += 1
        continue

    # Code block
    if stripped.startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_lines_buf = []
        else:
            in_code_block = False
            if code_lines_buf:
                add_code_block(code_lines_buf)
                code_lines_buf = []
        i += 1
        continue
    if in_code_block:
        code_lines_buf.append(line)
        i += 1
        continue

    # Markdown table
    if stripped.startswith('|') and '|' in stripped[1:]:
        if not in_table:
            in_table = True
            table_lines_buf = []
        table_lines_buf.append(stripped)
        i += 1
        continue
    else:
        if in_table:
            in_table = False
            add_table_from_md(table_lines_buf)
            table_lines_buf = []

    # HTML table — bỏ qua (chỉ lấy text trong <td>/<th>)
    if stripped.startswith('<table') or stripped.startswith('<tr') or \
       stripped.startswith('<thead') or stripped.startswith('<tbody') or \
       stripped.startswith('<colgroup') or stripped.startswith('<col ') or \
       stripped.startswith('</'):
        # Lấy text trong <td> hoặc <th>
        cell_matches = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', stripped, re.DOTALL)
        for cm in cell_matches:
            clean_cm = re.sub(r'<[^>]+>', '', cm).strip()
            if clean_cm:
                add_para(clean_cm, indent_cm=1.5, size=12)
        # <li>
        li_matches = re.findall(r'<li[^>]*><p>(.*?)</p>', stripped, re.DOTALL)
        for lm in li_matches:
            clean_lm = re.sub(r'<[^>]+>', '', lm).strip()
            if clean_lm:
                add_para('• ' + clean_lm, indent_cm=2.0, size=13)
        i += 1
        continue

    # Dòng phân cách ---
    if re.match(r'^-{3,}$', stripped):
        p = doc.add_paragraph()
        set_para_spacing(p, 3, 3, 1.0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # Headings
    h_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
    if h_match:
        level = len(h_match.group(1))
        heading_text = h_match.group(2)
        add_heading(heading_text, level)
        i += 1
        continue

    # Blockquote >
    if stripped.startswith('>'):
        content = stripped.lstrip('>').strip()
        if content:
            add_para(content, indent_cm=1.0, size=13,
                     space_before=2, space_after=2,
                     color=(80, 80, 80))
        i += 1
        continue

    # Danh sách có số: "1. ", "2. "...
    num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
    if num_match:
        num   = num_match.group(1)
        content_text = num_match.group(2)
        add_para(f'{num}.  {content_text}', indent_cm=0.5, size=13,
                 space_before=4, space_after=2)
        i += 1
        continue

    # Danh sách bullet: - hoặc *
    bullet_match = re.match(r'^[-*+]\s+(.+)', stripped)
    if bullet_match:
        add_para('•  ' + bullet_match.group(1), indent_cm=1.0, size=13,
                 space_before=2, space_after=2)
        i += 1
        continue

    # Danh sách bullet lồng nhau: "  - "
    sub_bullet = re.match(r'^\s{2,}[-*+]\s+(.+)', line)
    if sub_bullet:
        add_para('   –  ' + sub_bullet.group(1), indent_cm=1.5, size=13,
                 space_before=1, space_after=1)
        i += 1
        continue

    # Dòng trống
    if not stripped:
        i += 1
        continue

    # Đoạn văn bình thường
    add_para(stripped, size=13, space_before=0, space_after=6)
    i += 1

# Xử lý table còn lại
if in_table and table_lines_buf:
    add_table_from_md(table_lines_buf)

# ── Lưu file ──────────────────────────────────────────────────────────────────
out_path = "BaoCao_Hanzii.docx"
doc.save(out_path)
import sys
sys.stdout.buffer.write(("DONE: " + out_path + "\n").encode("utf-8"))
sys.stdout.buffer.write(("Tong so doan: " + str(len(doc.paragraphs)) + "\n").encode("utf-8"))
